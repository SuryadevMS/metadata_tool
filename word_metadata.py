from tkinter import filedialog, messagebox
import tkinter as tk
from docx import Document

class WordMetadataTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Word Metadata Tool")
        self.root.geometry("1000x720")
        
        self.selected_file = None

        self.file_label = tk.Label(root, text="Selected File:")
        self.file_label.pack(pady=10)

        self.file_entry = tk.Entry(root, width=50, state="disabled")
        self.file_entry.pack(pady=5)

        self.browse_button = tk.Button(root, text="Browse", command=self.browse_file)
        self.browse_button.pack(pady=5)

        self.extract_button = tk.Button(root, text="Extract Metadata", command=self.extract_metadata)
        self.extract_button.pack(pady=5)

        self.remove_button = tk.Button(root, text="Remove Metadata", command=self.remove_metadata)
        self.remove_button.pack(pady=5)

        self.export_button = tk.Button(root, text="Export Metadata", command=self.export_metadata)
        self.export_button.pack(pady=5)

        self.metadata_display = tk.Text(root, height=20, width=80)
        self.metadata_display.pack(pady=10)

    def browse_file(self):
        filetypes = (("Word files", "*.docx"), ("All files", "*.*"))
        file_path = filedialog.askopenfilename(title="Select Word File", filetypes=filetypes)
        if file_path:
            self.selected_file = file_path
            self.file_entry.config(state="normal")
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(state="disabled")

    def extract_metadata(self):
        if self.selected_file:
            try:
                document = Document(self.selected_file)
                core_properties = document.core_properties
                metadata_str = ""
                for prop in core_properties:
                    metadata_str += f"{prop}: {getattr(core_properties, prop)}\n"
                self.metadata_display.config(state='normal')
                self.metadata_display.delete("1.0", tk.END)
                self.metadata_display.insert(tk.END, metadata_str)
                self.metadata_display.config(state='disabled')
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {str(e)}")
        else:
            messagebox.showwarning("Word Metadata", "Please select a Word file.")

    def remove_metadata(self):
        if self.selected_file:
            try:
                document = Document(self.selected_file)
                for prop in document.core_properties:
                    setattr(document.core_properties, prop, None)
                save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx"), ("All files", "*.*")], title="Save Document Without Metadata")
                if save_path:
                    document.save(save_path)
                    messagebox.showinfo("Remove Metadata", "Metadata removed and document saved successfully.")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {str(e)}")
        else:
            messagebox.showwarning("Remove Metadata", "Please select a Word file.")

    def export_metadata(self):
        if self.selected_file:
            try:
                document = Document(self.selected_file)
                core_properties = document.core_properties
                save_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt"), ("All files", "*.*")], title="Save Metadata")
                if save_path:
                    with open(save_path, 'w') as f:
                        for prop in core_properties:
                            f.write(f"{prop}: {getattr(core_properties, prop)}\n")
                    messagebox.showinfo("Export Metadata", "Metadata exported successfully.")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {str(e)}")
        else:
            messagebox.showwarning("Export Metadata", "Please select a Word file.")

if __name__ == "__main__":
    root = tk.Tk()
    app = WordMetadataTool(root)
    root.mainloop()
